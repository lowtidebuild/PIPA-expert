from __future__ import annotations

import json
from zipfile import ZipFile

from docx import Document

from scripts.audit_document import extract_docx_text, prepare_audit_document


def _aggregated(text: str, start: int, end: int, label: str = "unknown") -> dict:
    return {
        "aggregated": [
            {
                "claim": {
                    "text": text,
                    "sentence_span": {"start": start, "end": end},
                    "claim_type": "citation",
                },
                "verdict": {
                    "claim": {
                        "text": text,
                        "sentence_span": {"start": start, "end": end},
                        "claim_type": "citation",
                    },
                    "label": label,
                    "verifier_name": "fixture",
                    "authority": 1.0,
                    "rationale": "fixture",
                    "evidence": [{"title": "source", "url": "https://example.test/source"}],
                },
            }
        ]
    }


def test_extract_docx_text_reads_paragraphs_and_tables(tmp_path) -> None:
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("본문 문단")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    doc.save(path)

    text = extract_docx_text(path)

    assert "본문 문단" in text
    assert "A | B" in text


def test_prepare_docx_without_aggregated_writes_source_and_skipped_status(tmp_path) -> None:
    path = tmp_path / "opinion.docx"
    doc = Document()
    doc.add_paragraph("개인정보 보호법 제15조 검토")
    doc.save(path)

    outputs = prepare_audit_document(path, tmp_path / "audit")

    assert outputs["source_md"].read_text(encoding="utf-8").strip() == "개인정보 보호법 제15조 검토"
    status = json.loads(outputs["audit_status"].read_text(encoding="utf-8"))
    assert status["status"] == "skipped"
    assert "aggregated verdict JSON not provided" in status["reason"]
    assert outputs["audited_md"] is None


def test_prepare_markdown_with_aggregated_renders_audited_sidecar(tmp_path) -> None:
    sentence = "개인정보 보호법 제15조는 동의 요건을 규정한다."
    md = tmp_path / "opinion.md"
    md.write_text(sentence + "\n", encoding="utf-8")
    aggregated = tmp_path / "aggregated.json"
    aggregated.write_text(json.dumps(_aggregated(sentence, 0, len(sentence)), ensure_ascii=False), encoding="utf-8")

    outputs = prepare_audit_document(md, tmp_path / "audit", aggregated_path=aggregated)

    rendered = outputs["audited_md"].read_text(encoding="utf-8")
    assert "[Partially Unverified]" in rendered
    assert "Citation Audit Log" in rendered
    status = json.loads(outputs["audit_status"].read_text(encoding="utf-8"))
    assert status["status"] == "complete"


def test_prepare_docx_with_aggregated_can_append_docx_log(tmp_path) -> None:
    sentence = "개인정보 보호법 제15조는 동의 요건을 규정한다."
    path = tmp_path / "opinion.docx"
    doc = Document()
    doc.add_paragraph(sentence)
    doc.save(path)
    aggregated = tmp_path / "aggregated.json"
    aggregated.write_text(json.dumps(_aggregated(sentence, 0, len(sentence)), ensure_ascii=False), encoding="utf-8")

    outputs = prepare_audit_document(path, tmp_path / "audit", aggregated_path=aggregated, append_docx=True)

    assert outputs["audited_docx"].exists()
    xml = ZipFile(outputs["audited_docx"]).read("word/document.xml").decode("utf-8")
    assert "Citation Audit Log" in xml
    assert "Citation audit status: complete" in xml
