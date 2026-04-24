from __future__ import annotations

from docx import Document

from scripts.validate_opinion_artifact import validate_artifact


VALID_TEXT = """# 개인정보보호법 검토 메모

[AI 생성 고지] 본 분석 메모는 AI 시스템의 지원을 받아 작성되었습니다.

이 문서는 법률 자문이 아니며, 구체적 사안은 전문가 검토가 필요합니다.

Citation audit status: complete

[VERIFIED] [Grade A] 개인정보 보호법 제15조
"""


def test_validate_artifact_accepts_valid_markdown(tmp_path) -> None:
    path = tmp_path / "valid.md"
    path.write_text(VALID_TEXT, encoding="utf-8")

    result = validate_artifact(path, require_audit=True)

    assert result.ok is True
    assert result.issues == []


def test_validate_artifact_rejects_placeholder(tmp_path) -> None:
    path = tmp_path / "placeholder.md"
    path.write_text(VALID_TEXT + "\n[수신인 이름]\n", encoding="utf-8")

    result = validate_artifact(path, require_audit=True)

    assert result.ok is False
    assert any(issue.code == "placeholder" for issue in result.issues)


def test_validate_artifact_rejects_missing_audit_when_required(tmp_path) -> None:
    path = tmp_path / "no-audit.md"
    path.write_text(VALID_TEXT.replace("Citation audit status: complete\n\n", ""), encoding="utf-8")

    result = validate_artifact(path, require_audit=True)

    assert result.ok is False
    assert any(issue.code == "missing_audit_status" for issue in result.issues)


def test_validate_artifact_reads_docx_text(tmp_path) -> None:
    path = tmp_path / "valid.docx"
    doc = Document()
    doc.add_paragraph("[AI 생성 고지] 본 분석 메모는 AI 시스템의 지원을 받아 작성되었습니다.")
    doc.add_paragraph("이 문서는 법률 자문이 아니며, 구체적 사안은 전문가 검토가 필요합니다.")
    doc.add_paragraph("Citation audit status: complete")
    doc.save(path)

    result = validate_artifact(path, require_audit=True)

    assert result.ok is True
