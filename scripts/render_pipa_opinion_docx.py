#!/usr/bin/env python3
"""Render structured PIPA opinion artifacts to DOCX and Markdown copies."""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from scripts.docx_citation_appendix import append_citation_audit_log, inject_unverified_tags, load_aggregated
from scripts.lib.audit_status import load_json, write_json
from scripts.lib.opinion_model import OpinionArtifact, load_opinion_artifact
from scripts.lib.paths import output_dir
from scripts.render_audit_append import render_audit_append


FONT_LATIN = "Times New Roman"
FONT_CJK = "맑은 고딕"
NAVY = RGBColor(0x1B, 0x2A, 0x4A)


@dataclass(frozen=True)
class RenderedOpinion:
    docx_path: Path
    markdown_path: Path
    audit_status_path: Path | None = None


def render_opinion_artifact(
    input_path: str | Path,
    *,
    out: str | Path | None = None,
    aggregated_path: str | Path | None = None,
    audit_status_path: str | Path | None = None,
) -> RenderedOpinion:
    """Render an OpinionArtifact JSON file into deterministic DOCX + Markdown."""
    source = Path(input_path)
    artifact = load_opinion_artifact(source)
    docx_path, markdown_path = _output_paths(source, out)
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    source_md = artifact.to_markdown()
    body_md = source_md
    markdown_copy = source_md
    aggregated: dict[str, Any] | None = None
    audit_status: dict[str, Any] | None = None
    status_output: Path | None = Path(audit_status_path) if audit_status_path else None

    if aggregated_path:
        aggregated = load_aggregated(aggregated_path)
        base_status = load_json(audit_status_path) if audit_status_path and Path(audit_status_path).exists() else None
        markdown_copy, audit_status = render_audit_append(source_md, aggregated, base_status=base_status)
        injection = inject_unverified_tags(source_md, aggregated, return_result=True)
        body_md = injection.body_md
        if status_output is None:
            status_output = markdown_path.with_suffix(".audit_status.json")
        write_json(status_output, audit_status)

    markdown_path.write_text(markdown_copy, encoding="utf-8")
    doc = _build_docx(body_md)
    if aggregated is not None:
        append_citation_audit_log(doc, aggregated, audit_status=audit_status)
    doc.save(docx_path)
    return RenderedOpinion(docx_path=docx_path, markdown_path=markdown_path, audit_status_path=status_output)


def _output_paths(input_path: Path, out: str | Path | None) -> tuple[Path, Path]:
    if out is None:
        out_dir = output_dir()
        stem = _safe_stem(input_path.stem)
        return out_dir / f"{stem}.docx", out_dir / f"{stem}.md"

    out_path = Path(out)
    if out_path.suffix.lower() == ".docx":
        return out_path, out_path.with_suffix(".md")

    stem = _safe_stem(input_path.stem)
    return out_path / f"{stem}.docx", out_path / f"{stem}.md"


def _safe_stem(value: str) -> str:
    stem = re.sub(r"[^0-9A-Za-z가-힣._-]+", "-", value).strip("-._")
    return stem or "opinion"


def _build_docx(markdown_text: str) -> Document:
    doc = Document()
    _configure_document(doc)
    _render_markdown(doc, markdown_text)
    return doc


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(11)
    _set_style_cjk(normal)

    for style_name, size in [("Title", 16), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = NAVY
        _set_style_cjk(style)


def _set_style_cjk(style: Any) -> None:
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)


def _render_markdown(doc: Document, markdown_text: str) -> None:
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_run(paragraph, line[2:].strip(), bold=True, size=16, color=NAVY)
        elif line.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 2")
            _add_run(paragraph, line[3:].strip(), bold=True, size=13, color=NAVY)
        elif line.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 3")
            _add_run(paragraph, line[4:].strip(), bold=True, size=11, color=NAVY)
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_run(paragraph, line[2:].strip())
        elif line.lstrip().startswith(">"):
            text = line.lstrip()[1:].strip()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1.0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            _add_run(paragraph, text)
        else:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1.2
            _add_inline_runs(paragraph, line)


def _add_inline_runs(paragraph: Any, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _add_run(paragraph, part[2:-2], bold=True)
        else:
            _add_run(paragraph, part)


def _add_run(
    paragraph: Any,
    text: str,
    *,
    bold: bool = False,
    size: int = 11,
    color: RGBColor | None = None,
) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = FONT_LATIN
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    _set_run_cjk(run)


def _set_run_cjk(run: Any) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_json", help="OpinionArtifact JSON input.")
    parser.add_argument("--out", help="Output directory or .docx path. Defaults to PIPA_OUTPUT_DIR.")
    parser.add_argument("--aggregated", help="Optional citation-auditor aggregated JSON.")
    parser.add_argument("--audit-status", help="Optional audit_status.json input/output path.")
    parser.add_argument("--json", action="store_true", help="Print rendered output paths as JSON.")
    args = parser.parse_args(argv)

    rendered = render_opinion_artifact(
        args.input_json,
        out=args.out,
        aggregated_path=args.aggregated,
        audit_status_path=args.audit_status,
    )
    payload = {
        "docx": str(rendered.docx_path),
        "markdown": str(rendered.markdown_path),
        "audit_status": str(rendered.audit_status_path) if rendered.audit_status_path else None,
    }
    if args.json:
        print(json.dumps(payload, ensure_ascii=False, indent=2))
    else:
        print(f"DOCX: {rendered.docx_path}")
        print(f"Markdown: {rendered.markdown_path}")
        if rendered.audit_status_path:
            print(f"Audit status: {rendered.audit_status_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
