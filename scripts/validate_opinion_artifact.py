"""Validate generated PIPA opinion/memo artifacts before delivery."""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import asdict, dataclass, field
from pathlib import Path

from docx import Document


PLACEHOLDER_PATTERNS = [
    re.compile(pattern)
    for pattern in [
        r"\[수신인 이름\]",
        r"\[수신인 직위\]",
        r"\[회사/기관명\]",
        r"\[주소\]",
        r"\[번호\]",
        r"\[이메일\]",
        r"\[웹사이트\]",
        r"\[직위\]",
        r"\[분석 내용[^\]]*\]",
        r"\[쟁점\s*\d+[^\]]*\]",
        r"\[쟁점\s*\d+에 대한 결론\]",
        r"제XX조",
        r"제2025-XX호",
        r"# \.\.\.",
    ]
]


@dataclass
class ValidationIssue:
    code: str
    message: str


@dataclass
class ValidationResult:
    path: str
    ok: bool
    issues: list[ValidationIssue] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "path": self.path,
            "ok": self.ok,
            "issues": [asdict(issue) for issue in self.issues],
        }


def validate_artifact(
    path: str | Path,
    *,
    require_audit: bool = False,
    audit_status_path: str | Path | None = None,
) -> ValidationResult:
    artifact_path = Path(path)
    text = extract_text(artifact_path)
    issues: list[ValidationIssue] = []

    for pattern in PLACEHOLDER_PATTERNS:
        match = pattern.search(text)
        if match:
            issues.append(
                ValidationIssue(
                    code="placeholder",
                    message=f"Placeholder-like text remains: {match.group(0)}",
                )
            )

    if "AI 생성" not in text and "AI system" not in text and "AI-generated" not in text:
        issues.append(ValidationIssue(code="missing_ai_notice", message="AI generation notice is missing."))

    if "법률 자문" not in text and "legal advice" not in text.lower():
        issues.append(ValidationIssue(code="missing_disclaimer", message="Legal-advice disclaimer is missing."))

    audit_status = _load_audit_status(audit_status_path)
    if require_audit and not _has_audit_marker(text, audit_status):
        issues.append(
            ValidationIssue(
                code="missing_audit_status",
                message="Citation audit status or appendix is required but missing.",
            )
        )
    if audit_status and audit_status.get("status") in {"partial", "failed", "skipped"}:
        reason = audit_status.get("reason") or audit_status.get("status")
        if str(reason) not in text:
            issues.append(
                ValidationIssue(
                    code="audit_status_not_disclosed",
                    message=f"Audit status is {audit_status.get('status')} but not disclosed in artifact text.",
                )
            )

    return ValidationResult(path=str(artifact_path), ok=not issues, issues=issues)


def extract_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        doc = Document(path)
        parts = [paragraph.text for paragraph in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    return path.read_text(encoding="utf-8")


def _load_audit_status(path: str | Path | None) -> dict | None:
    if path is None:
        return None
    status_path = Path(path)
    if not status_path.exists():
        return None
    return json.loads(status_path.read_text(encoding="utf-8"))


def _has_audit_marker(text: str, audit_status: dict | None) -> bool:
    if "Citation Audit Log" in text or "Citation audit status" in text:
        return True
    return bool(audit_status and audit_status.get("status"))


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("path", help="Markdown or DOCX artifact path.")
    parser.add_argument("--require-audit", action="store_true", help="Require citation audit status/appendix.")
    parser.add_argument("--audit-status", help="Optional audit_status.json sidecar.")
    parser.add_argument("--json", action="store_true", help="Print machine-readable JSON.")
    args = parser.parse_args(argv)

    result = validate_artifact(args.path, require_audit=args.require_audit, audit_status_path=args.audit_status)
    if args.json:
        print(json.dumps(result.to_dict(), ensure_ascii=False, indent=2))
    else:
        if result.ok:
            print(f"ok: {result.path}")
        else:
            print(f"FAIL: {result.path}")
            for issue in result.issues:
                print(f"- {issue.code}: {issue.message}")
    return 0 if result.ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
