"""DOCX adapter for citation-auditor aggregated verdict JSON.

The citation-auditor package is markdown-native. This module is the PIPA
project-local bridge for DOCX opinions: it tags failing claims before markdown
content is embedded in a DOCX renderer, and appends a compact audit-log
appendix to the python-docx document before save.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


FONT_LATIN = "Times New Roman"
FONT_CJK = "맑은 고딕"
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GRAY = RGBColor(0x59, 0x59, 0x59)
LIGHT_FILL = "EAF2F8"

_TAG_FOR_LABEL = {
    "contradicted": "[Unverified]",
    "unknown": "[Partially Unverified]",
}

_APPENDIX_HEADING = "부록: 검증 로그 (Citation Audit Log)"
_APPENDIX_NOTE = (
    "본 부록은 본 산출물에 포함된 사실·인용 주장에 대한 자동 citation audit 결과입니다. "
    "자동 감사는 전문가 검토를 대체하지 않으며, Contradicted 또는 Unknown 항목은 "
    "최종 검토 시 독립 확인이 필요합니다."
)
_HEADERS = ["#", "클레임 (Claim)", "판정 (Verdict)", "Verifier", "근거 (Evidence)"]


def load_aggregated(path: str | Path) -> dict[str, Any]:
    """Load an aggregated citation-auditor JSON file."""
    return json.loads(Path(path).read_text(encoding="utf-8"))


def inject_unverified_tags(body_md: str, aggregated: dict[str, Any]) -> str:
    """Insert failing-verdict tags after claim sentence spans in markdown.

    Verified claims are left untouched. Spans are expected to be
    document-relative, as produced by `python -m citation_auditor aggregate`.
    Invalid spans are skipped rather than guessed.
    """
    insertions: list[tuple[int, str]] = []
    for item in _verdict_items(aggregated):
        verdict = item.get("verdict") or {}
        tag = _TAG_FOR_LABEL.get(str(verdict.get("label", "")))
        if tag is None:
            continue
        claim = verdict.get("claim") or item.get("claim") or {}
        span = claim.get("sentence_span") or {}
        end = span.get("end")
        if not isinstance(end, int) or end < 0 or end > len(body_md):
            continue
        insertions.append((end, f" {tag}"))

    rendered = body_md
    for offset, tag in sorted(set(insertions), key=lambda value: value[0], reverse=True):
        rendered = rendered[:offset] + tag + rendered[offset:]
    return rendered


def append_citation_audit_log(doc: Any, aggregated: dict[str, Any]) -> None:
    """Append a bilingual citation audit appendix to a python-docx Document."""
    doc.add_page_break()
    _add_heading(doc, _APPENDIX_HEADING)
    _add_small_paragraph(doc, _APPENDIX_NOTE)

    items = _verdict_items(aggregated)
    if not items:
        _add_small_paragraph(
            doc,
            "[Citation Audit Skipped] 추출된 검증 대상 클레임이 없거나 verifier 결과가 없습니다.",
            color=RGBColor(0x80, 0x00, 0x00),
        )
        return

    table = doc.add_table(rows=1, cols=len(_HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(_HEADERS):
        cell = table.rows[0].cells[index]
        _shade_cell(cell, LIGHT_FILL)
        _set_cell_text(cell, header, bold=True)

    for row_index, item in enumerate(items, start=1):
        verdict = item.get("verdict") or {}
        claim = verdict.get("claim") or item.get("claim") or {}
        row = table.add_row().cells
        _set_cell_text(row[0], str(row_index))
        _set_cell_text(row[1], _truncate(str(claim.get("text", "")), 140))
        _set_cell_text(row[2], _format_label(str(verdict.get("label", "unknown"))))
        _set_cell_text(row[3], str(verdict.get("verifier_name", "-")))
        _set_cell_text(row[4], _format_evidence(verdict.get("evidence") or []))


def _verdict_items(aggregated: dict[str, Any]) -> list[dict[str, Any]]:
    items = aggregated.get("aggregated")
    return items if isinstance(items, list) else []


def _add_heading(doc: Any, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY
    _set_run_fonts(run)


def _add_small_paragraph(doc: Any, text: str, color: RGBColor = GRAY) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = color
    _set_run_fonts(run)


def _set_cell_text(cell: Any, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.5)
    _set_run_fonts(run)


def _set_run_fonts(run: Any) -> None:
    run.font.name = FONT_LATIN
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.xpath("./w:shd")
    if shd:
        shd[0].set(qn("w:fill"), fill)
        return
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _format_label(label: str) -> str:
    if label == "verified":
        return "Verified"
    if label == "contradicted":
        return "Contradicted"
    return "Unknown"


def _format_evidence(evidence: list[dict[str, Any]]) -> str:
    parts: list[str] = []
    for item in evidence:
        url = str(item.get("url", "")).strip()
        title = str(item.get("title", "")).strip()
        if title and url and title != url:
            parts.append(f"{title} ({url})")
        elif url:
            parts.append(url)
        elif title:
            parts.append(title)
    return _truncate("; ".join(parts), 100) if parts else "-"


def _truncate(text: str, max_len: int) -> str:
    collapsed = " ".join(text.split())
    if len(collapsed) <= max_len:
        return collapsed
    return collapsed[: max_len - 1] + "…"
