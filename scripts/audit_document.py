#!/usr/bin/env python3
"""Prepare or render citation audits for Markdown and DOCX artifacts."""

from __future__ import annotations

import argparse
import shutil
from pathlib import Path
import sys

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from scripts.docx_citation_appendix import append_citation_audit_log, load_aggregated
from scripts.lib.audit_status import build_audit_status, write_json
from scripts.render_audit_append import render_audit_append


def extract_docx_text(path: str | Path) -> str:
    """Extract simple paragraph/table text from a DOCX without external converters."""
    doc = Document(path)
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip() + "\n"


def prepare_audit_document(
    input_path: str | Path,
    out_dir: str | Path,
    *,
    aggregated_path: str | Path | None = None,
    append_docx: bool = False,
) -> dict[str, Path | None]:
    """Prepare sidecar files and optionally render audit annotations."""
    source = Path(input_path)
    output_dir = Path(out_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    source_md = _source_markdown_path(source, output_dir)
    status_path = output_dir / f"{source.stem}.audit_status.json"
    audited_md = output_dir / f"{source.stem}.audited.md"
    audited_docx: Path | None = None

    if source.suffix.lower() == ".docx":
        source_md.write_text(extract_docx_text(source), encoding="utf-8")
    elif source.suffix.lower() in {".md", ".markdown", ".txt"}:
        source_md.write_text(source.read_text(encoding="utf-8"), encoding="utf-8")
    else:
        source_md.write_text(source.read_text(encoding="utf-8"), encoding="utf-8")

    if aggregated_path is None:
        status = build_audit_status(
            {"aggregated": []},
            input_path=str(source),
            reason="aggregated verdict JSON not provided; source markdown sidecar prepared only",
        )
        write_json(status_path, status)
        return {
            "source_md": source_md,
            "audited_md": None,
            "audit_status": status_path,
            "audited_docx": None,
        }

    aggregated = load_aggregated(aggregated_path)
    rendered, status = render_audit_append(source_md.read_text(encoding="utf-8"), aggregated)
    audited_md.write_text(rendered, encoding="utf-8")
    status["input_path"] = str(source)
    write_json(status_path, status)

    if source.suffix.lower() == ".docx" and append_docx:
        audited_docx = output_dir / f"{source.stem}.audited.docx"
        shutil.copyfile(source, audited_docx)
        doc = Document(audited_docx)
        append_citation_audit_log(doc, aggregated, audit_status=status)
        doc.save(audited_docx)

    return {
        "source_md": source_md,
        "audited_md": audited_md,
        "audit_status": status_path,
        "audited_docx": audited_docx,
    }


def _source_markdown_path(source: Path, output_dir: Path) -> Path:
    if source.suffix.lower() in {".md", ".markdown"}:
        return output_dir / f"{source.stem}.audit-source.md"
    return output_dir / f"{source.stem}.audit-source.md"


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", help="Markdown, text, or DOCX artifact.")
    parser.add_argument("--out", required=True, help="Output directory for audit sidecars.")
    parser.add_argument("--aggregated", help="Optional citation-auditor aggregated verdict JSON.")
    parser.add_argument("--append-docx", action="store_true", help="For DOCX input, write a copied DOCX with audit appendix.")
    args = parser.parse_args(argv)

    outputs = prepare_audit_document(
        args.input,
        args.out,
        aggregated_path=args.aggregated,
        append_docx=args.append_docx,
    )
    for name, path in outputs.items():
        if path is not None:
            print(f"{name}: {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
